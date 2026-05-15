import re
import os
import yaml
from typing import List, Optional, Dict, Any
from .base import BaseParser, ParsedDocument, DocumentSection, TermDefinition

try:
    from docx import Document
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocxParser(BaseParser):
    """Word (.docx) 文档解析器 - 支持通用标准和化妆品检验方法"""

    # GB/T 标准中术语章节的标志词
    TERM_SECTION_KEYWORDS = ["术语", "定义", "缩略语", "术语和定义", "术语、定义"]

    # 化妆品检验方法文档类型识别关键词
    INSPECTION_TYPE_KEYWORDS = [
        "检验方法", "Determination of", "Test method",
        "化妆品中", "牙膏中"
    ]

    # 检验方法特有章节名模式（用于增强标题识别）
    INSPECTION_SECTION_PATTERNS = [
        r"^方法提要",
        r"^方法原理",
        r"^试剂和材料",
        r"^试剂、材料",
        r"^试剂材料",
        r"^仪器和设备",
        r"^仪器、设备",
        r"^分析步骤",
        r"^操作步骤",
        r"^测定步骤",
        r"^分析结果的表述",
        r"^结果计算",
        r"^图谱$",
        r"^附录\s*[A-Z]",
        r"^附件\s*\d+",
    ]

    def __init__(self):
        """初始化，加载检验方法规则"""
        self._inspection_rules = None
        self._load_inspection_rules()

    def _load_inspection_rules(self):
        """加载检验方法规则配置"""
        try:
            rules_path = os.path.join(
                os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                "rules", "inspection_method.yaml"
            )
            if os.path.exists(rules_path):
                with open(rules_path, "r", encoding="utf-8") as f:
                    self._inspection_rules = yaml.safe_load(f)
        except Exception:
            self._inspection_rules = None

    def parse(self, file_path: str) -> ParsedDocument:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，请运行: pip install python-docx")

        doc = Document(file_path)
        title = self._extract_title(doc)
        sections, full_text = self._extract_sections(doc)
        terms = self._extract_terms(sections)
        tables = self._extract_tables(doc)

        # 自动识别文档类型
        doc_type = self._detect_doc_type(title, full_text)

        # 提取附录A标准品信息
        appendix_data = {}
        if doc_type == "inspection_method":
            appendix_data = self._extract_appendix_data(doc)

        metadata = {
            "file_type": "docx",
            "paragraphs": len(doc.paragraphs),
            "doc_type": doc_type,
            "tables": tables,
            "appendix_data": appendix_data,
        }

        return ParsedDocument(
            title=title,
            sections=sections,
            terms=terms,
            full_text=full_text,
            metadata=metadata
        )

    def _detect_doc_type(self, title: str, full_text: str) -> str:
        """自动识别文档类型"""
        # 优先从标题判断
        check_text = title
        for keyword in self.INSPECTION_TYPE_KEYWORDS:
            if keyword in check_text:
                return "inspection_method"

        # 标题未匹配时检查正文前500字
        check_text = full_text[:500]
        for keyword in self.INSPECTION_TYPE_KEYWORDS[:2]:  # 只用最关键的词
            if keyword in check_text:
                return "inspection_method"

        return "general_standard"

    def _extract_title(self, doc) -> str:
        """提取文档标题（通常是第一个非空段落或 Title 样式段落）"""
        for para in doc.paragraphs:
            if para.style.name in ("Title", "标题") and para.text.strip():
                return para.text.strip()
        # 回退：取第一个较长的非空行
        for para in doc.paragraphs:
            t = para.text.strip()
            if t and len(t) > 5:
                return t
        return "未知标题"

    def _extract_sections(self, doc) -> tuple:
        """提取章节结构和全文（段落+表格内容）"""
        sections: List[DocumentSection] = []
        full_lines = []
        current_section: Optional[DocumentSection] = None
        current_content_lines = []

        # 预扫描判断是否为检验方法文档
        is_inspection = False
        for para in doc.paragraphs[:20]:
            t = para.text.strip()
            if t:
                for keyword in self.INSPECTION_TYPE_KEYWORDS:
                    if keyword in t:
                        is_inspection = True
                        break
                if is_inspection:
                    break

        # 遍历 body 下的所有子元素（段落和表格按文档顺序）
        W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        body = doc.element.body

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":  # 段落
                from docx.text.paragraph import Paragraph
                para = Paragraph(child, doc)
                text = para.text.strip()
                if not text:
                    continue

                full_lines.append(text)
                style_name = para.style.name.lower()
                is_heading = self._is_heading_paragraph(text, style_name, is_inspection)

                if is_heading:
                    if current_section is not None:
                        current_section.content = "\n".join(current_content_lines)
                        sections.append(current_section)
                        current_content_lines = []

                    number, title_text, level = self._parse_heading(text, is_inspection)
                    current_section = DocumentSection(
                        number=number,
                        title=title_text,
                        level=level,
                        content="",
                        raw_heading=text
                    )
                else:
                    if current_section is not None:
                        current_content_lines.append(text)

            elif tag == "tbl":  # 表格：将表格转为文本加入当前章节
                from docx.table import Table
                table = Table(child, doc)
                table_lines = []
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    table_lines.append(" | ".join(cells))
                table_text = "\n".join(table_lines)
                if table_text and current_section is not None:
                    current_content_lines.append("[表格]\n" + table_text + "\n[/表格]")
                full_lines.append(table_text)

        # 保存最后一节
        if current_section is not None:
            current_section.content = "\n".join(current_content_lines)
            sections.append(current_section)

        return sections, "\n".join(full_lines)

    def _is_heading_paragraph(self, text: str, style_name: str, is_inspection: bool) -> bool:
        """判断段落是否为标题段落"""
        # 样式判断
        if "heading" in style_name or "标题" in style_name:
            return True

        # 通用编号模式：1, 1.1, A.1 等
        if re.match(r'^(\d+|[A-Z]\.\d+)(\.\d+)*\s+\S', text):
            return True

        # 检验方法特有标题模式
        if is_inspection:
            for pattern in self.INSPECTION_SECTION_PATTERNS:
                if re.match(pattern, text):
                    return True
            # 匹配 "附件X" 格式
            if re.match(r'^附件\s*\d+', text):
                return True

        return False

    def _parse_heading(self, text: str, is_inspection: bool) -> tuple:
        """解析标题，返回 (编号, 标题文本, 层级)"""
        # 尝试提取编号
        number = self._extract_section_number(text)

        if number:
            level = self._get_level_from_number(number)
            # 去除编号得到标题文本
            title_text = re.sub(r'^[A-Z]?\d+(?:\.\d+)*\s+', '', text).strip()
        else:
            # 无编号标题
            level = 0
            title_text = text

        # 检验方法特有标题的编号推断
        if is_inspection and not number:
            inferred = self._infer_inspection_number(title_text)
            if inferred:
                number, level = inferred

        return number, title_text, level

    def _infer_inspection_number(self, title_text: str) -> Optional[tuple]:
        """根据检验方法章节名推断编号"""
        # 这些章节通常没有显式编号前缀，但有固定层级
        known_chapters = {
            "范围": ("1", 1),
            "方法提要": ("2", 1),
            "方法原理": ("2", 1),
            "试剂和材料": ("3", 1),
            "试剂、材料": ("3", 1),
            "仪器和设备": ("4", 1),
            "仪器、设备": ("4", 1),
            "分析步骤": ("5", 1),
            "操作步骤": ("5", 1),
            "测定步骤": ("5", 1),
            "计算": ("6", 1),
            "结果计算": ("6", 1),
            "分析结果的表述": ("6", 1),
            "图谱": ("7", 1),
        }
        return known_chapters.get(title_text)

    def _extract_tables(self, doc) -> List[Dict[str, Any]]:
        """提取文档中的所有表格"""
        tables = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            tables.append({
                "index": i,
                "rows": rows,
                "row_count": len(rows),
                "col_count": len(rows[0]) if rows else 0,
                "header": rows[0] if rows else [],
            })
        return tables

    def _extract_appendix_data(self, doc) -> Dict[str, Any]:
        """从文档中提取附录信息（标准品信息表等）"""
        result = {}

        # 查找附录A相关表格
        appendix_a_tables = []
        in_appendix_a = False

        for para in doc.paragraphs:
            text = para.text.strip()
            if re.match(r'^附录\s*A', text, re.IGNORECASE):
                in_appendix_a = True
                continue
            # 遇到下一个附录时停止
            if in_appendix_a and re.match(r'^附录\s*[B-Z]', text, re.IGNORECASE):
                in_appendix_a = False
                break

        # 从所有表格中查找标准品信息表（通过列名判断）
        tables = self._extract_tables(doc)
        standard_columns = ["中文名称", "英文名称", "CAS号", "序号", "分子式", "纯度"]

        for table_info in tables:
            header = table_info.get("header", [])
            match_count = sum(1 for col in standard_columns if any(col in h for h in header))
            if match_count >= 3:  # 至少匹配3个标准列名
                appendix_a_tables.append(table_info)

        if appendix_a_tables:
            result["standard_compound_tables"] = appendix_a_tables
            result["has_appendix_a"] = True
        else:
            result["has_appendix_a"] = False

        return result

    def _extract_terms(self, sections: List[DocumentSection]) -> List[TermDefinition]:
        """从术语章节提取术语定义"""
        terms = []
        term_sections = [
            s for s in sections
            if any(kw in s.title for kw in self.TERM_SECTION_KEYWORDS)
            or any(kw in s.raw_heading for kw in self.TERM_SECTION_KEYWORDS)
        ]

        for sec in term_sections:
            # 尝试解析 "X.X 术语名\n  定义文本" 模式
            lines = sec.content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                # 术语编号格式：3.1, 3.1.1 等
                m = re.match(r'^(\d+\.\d+(?:\.\d+)?)\s+(.+)', line)
                if m:
                    term_text = m.group(2).strip()
                    # 下一行作为定义
                    definition = ""
                    if i + 1 < len(lines):
                        definition = lines[i + 1].strip()
                    terms.append(TermDefinition(
                        term=term_text,
                        definition=definition,
                        section=sec.number
                    ))
                i += 1

        return terms
