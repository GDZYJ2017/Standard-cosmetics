import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from typing import Optional, List
import uuid
import logging

from models.database import get_db, Standard, ReviewTask, ReviewReport, ReviewBatch
from models.schemas import StandardOut, ReviewTaskOut, ReportOut, ApiResponse
from services.file_service import validate_upload, save_standard_file, delete_standard_file, save_draft_file
from services.report_service import ReportService

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api", tags=["standards"])
report_service = ReportService()


# 全局审查任务状态（简易存储，生产环境用 Redis）
_task_status: dict = {}


@router.post("/standards/upload")
async def upload_standard(
    file: UploadFile = File(...),
    name: str = Form(""),
    number: str = Form(""),
    description: str = Form("")
):
    """上传参考标准文件"""
    content = await file.read()
    ok, err = validate_upload(file.filename, len(content))
    if not ok:
        return ApiResponse(success=False, message=err)

    file_info = await save_standard_file(content, file.filename)

    # 保存到数据库
    async for db in get_db():
        std = Standard(
            id=file_info["id"],
            name=name or file.filename,
            number=number,
            file_name=file_info["file_name"],
            file_path=file_info["file_path"],
            file_type=file_info["file_type"],
            file_size=file_info["file_size"],
            description=description
        )
        db.add(std)
        await db.commit()
        await db.refresh(std)
        return ApiResponse(success=True, message="上传成功", data=StandardOut.model_validate(std).model_dump())


@router.get("/standards")
async def list_standards():
    """获取参考标准列表"""
    async for db in get_db():
        from sqlalchemy import select
        result = await db.execute(select(Standard).order_by(Standard.upload_time.desc()))
        standards = result.scalars().all()
        return ApiResponse(
            success=True,
            data=[StandardOut.model_validate(s).model_dump() for s in standards]
        )


@router.delete("/standards/{standard_id}")
async def delete_standard(standard_id: str):
    """删除参考标准"""
    async for db in get_db():
        from sqlalchemy import select
        result = await db.execute(select(Standard).where(Standard.id == standard_id))
        std = result.scalar_one_or_none()
        if not std:
            raise HTTPException(status_code=404, detail="标准不存在")
        delete_standard_file(std.file_path)
        await db.delete(std)
        await db.commit()
        return ApiResponse(success=True, message="删除成功")


async def _combine_standard_files(standard_paths: list, task_id: str) -> str:
    """将多个参考标准文件合并为一个临时 docx 文件"""
    import tempfile
    from docx import Document

    from config import UPLOAD_DIR
    combined_path = os.path.join(UPLOAD_DIR, f"combined_ref_{task_id}.docx")
    doc = Document()

    for i, path in enumerate(standard_paths):
        try:
            src = Document(path)
            # 添加分隔标题
            if i > 0:
                doc.add_heading(f"（参考标准 {i + 1}）", level=1)
            for para in src.paragraphs:
                text = para.text.strip()
                if text:
                    doc.add_paragraph(text)
            # 处理表格
            for table in src.tables:
                t = doc.add_table(rows=len(table.rows), cols=len(table.columns))
                for ri, row in enumerate(table.rows):
                    for ci, cell in enumerate(row.cells):
                        t.rows[ri].cells[ci].text = cell.text
        except Exception as e:
            logger.warning(f"合并标准文件失败 {path}: {e}")

    doc.save(combined_path)
    return combined_path


@router.post("/reviews/create")
async def create_review(
    name: str = Form(""),
    standard_ids: str = Form(""),
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = None
):
    """创建审查任务"""
    if not standard_ids:
        return ApiResponse(success=False, message="请选择参考标准")

    # 解析多个 standard_id（逗号分隔）
    std_id_list = [sid.strip() for sid in standard_ids.split(",") if sid.strip()]
    if not std_id_list:
        return ApiResponse(success=False, message="请选择参考标准")

    # 查找所有参考标准
    async for db in get_db():
        from sqlalchemy import select
        result = await db.execute(select(Standard).where(Standard.id.in_(std_id_list)))
        standards = result.scalars().all()
        if not standards:
            return ApiResponse(success=False, message="参考标准不存在")
        if len(standards) != len(std_id_list):
            found_ids = {s.id for s in standards}
            missing = [sid for sid in std_id_list if sid not in found_ids]
            return ApiResponse(success=False, message=f"部分参考标准不存在: {', '.join(missing)}")

        # 验证并保存初稿
        content = await file.read()
        ok, err = validate_upload(file.filename, len(content))
        if not ok:
            return ApiResponse(success=False, message=err)

        from services.file_service import save_draft_file
        file_info = await save_draft_file(content, file.filename)

        task_id = uuid.uuid4().hex

        # 收集所有参考标准路径
        std_paths = [s.file_path for s in standards]
        std_names = " / ".join([s.name for s in standards])
        # 单标准用 ref_path（兼容），多标准用 ref_paths（共识模式）
        ref_path = std_paths[0] if len(std_paths) == 1 else std_paths[0]
        ref_paths = std_paths if len(std_paths) > 1 else None

        task = ReviewTask(
            id=task_id,
            name=name or f"审查-{file.filename}",
            standard_id=standard_ids,  # 存储逗号分隔的ID
            draft_file_name=file_info["file_name"],
            draft_file_path=file_info["file_path"],
            draft_file_type=file_info["file_type"],
            status="pending",
            progress=0
        )
        db.add(task)
        await db.commit()
        await db.refresh(task)

        # 先立即更新状态为 running，让前端能感知到任务已启动
        from sqlalchemy import select
        result2 = await db.execute(select(ReviewTask).where(ReviewTask.id == task_id))
        task_now = result2.scalar_one_or_none()
        if task_now:
            task_now.status = "running"
            task_now.progress = 5
            task_now.current_step = "等待处理..."
            await db.commit()

        # 用 FastAPI BackgroundTasks 启动后台审查（多标准时启用共识合成模式）
        background_tasks.add_task(
            _run_review_task,
            task_id,
            ref_path,
            file_info["file_path"],
            task.name,
            ref_paths=ref_paths
        )

        return ApiResponse(success=True, message="审查任务已创建", data={
            "task_id": task_id,
            "status": "running"
        })


@router.post("/reviews/batch-create")
async def batch_create_review(
    batch_name: str = Form(""),
    standard_ids: str = Form(""),
    files: List[UploadFile] = File(...),
    background_tasks: BackgroundTasks = None
):
    """批量创建审查任务（单次上传多份初稿）"""
    if not standard_ids:
        return ApiResponse(success=False, message="请选择参考标准")
    if not files:
        return ApiResponse(success=False, message="请上传至少一份初稿")
    if len(files) > 20:
        return ApiResponse(success=False, message="单次批量任务不超过20份，请分批上传")

    std_id_list = [sid.strip() for sid in standard_ids.split(",") if sid.strip()]
    if not std_id_list:
        return ApiResponse(success=False, message="请选择参考标准")

    # 创建批次记录
    batch_id = uuid.uuid4().hex
    batch_display_name = batch_name.strip() or f"批量任务-{__import__('datetime').datetime.now().strftime('%Y%m%d%H%M')}"

    # 查询参考标准
    async for db in get_db():
        from sqlalchemy import select
        result = await db.execute(select(Standard).where(Standard.id.in_(std_id_list)))
        standards = result.scalars().all()
        if not standards:
            return ApiResponse(success=False, message="参考标准不存在")
        if len(standards) != len(std_id_list):
            found_ids = {s.id for s in standards}
            missing = [sid for sid in std_id_list if sid not in found_ids]
            return ApiResponse(success=False, message=f"部分参考标准不存在: {', '.join(missing)}")

        # 保存合并后的参考标准文件（所有标准合并为一份，供所有任务共用）
        import tempfile, os as _os
        with tempfile.TemporaryDirectory() as tmpdir:
            # 先保存各标准文件再合并
            std_paths = []
            for s in standards:
                if _os.path.exists(s.file_path):
                    std_paths.append(s.file_path)

            if len(std_paths) == 1:
                ref_path = std_paths[0]
                ref_paths = None
            else:
                # 多标准：传入完整路径列表（共识合成模式），不再合并文件
                ref_path = std_paths[0]
                ref_paths = std_paths

            # 创建批次记录
            batch = ReviewBatch(id=batch_id, name=batch_display_name, total_count=len(files))
            db.add(batch)

            task_ids = []
            for file in files:
                content = await file.read()
                ok, err = validate_upload(file.filename, len(content))
                if not ok:
                    continue  # 跳过无效文件

                file_info = await save_draft_file(content, file.filename)
                task_id = uuid.uuid4().hex
                task_name = batch_display_name + " - " + file.filename

                task = ReviewTask(
                    id=task_id,
                    batch_id=batch_id,
                    name=task_name,
                    standard_id=standard_ids,
                    draft_file_name=file_info["file_name"],
                    draft_file_path=file_info["file_path"],
                    draft_file_type=file_info["file_type"],
                    status="running",
                    progress=0
                )
                db.add(task)
                task_ids.append(task_id)

            await db.commit()

            # 启动所有审查任务（后台）
            for task_id in task_ids:
                # 查找该任务的draft_path（重新查一次，因为上面session已关闭）
                async for _db in get_db():
                    from sqlalchemy import select
                    res = await _db.execute(select(ReviewTask).where(ReviewTask.id == task_id))
                    t = res.scalar_one_or_none()
                    if t:
                        background_tasks.add_task(
                            _run_review_task,
                            task_id,
                            ref_path,
                            t.draft_file_path,
                            t.name,
                            ref_paths=ref_paths
                        )
                    break

            return ApiResponse(
                success=True,
                message=f"批量任务已创建，共 {len(task_ids)} 份",
                data={
                    "batch_id": batch_id,
                    "batch_name": batch_display_name,
                    "total_count": len(task_ids),
                    "task_ids": task_ids
                }
            )


async def _update_task_in_db(task_id: str, progress: int, step: str, status: str,
                             parsing_started_at=None, analysis_started_at=None,
                             score=None, critical_issues=0, major_issues=0, minor_issues=0):
    """在独立的 db session 中更新任务状态（含时间戳和报告摘要）"""
    from sqlalchemy import select
    try:
        async for db in get_db():
            result = await db.execute(select(ReviewTask).where(ReviewTask.id == task_id))
            task = result.scalar_one_or_none()
            if task:
                task.progress = progress
                task.current_step = step
                task.status = status
                if parsing_started_at:
                    task.parsing_started_at = parsing_started_at
                if analysis_started_at:
                    task.analysis_started_at = analysis_started_at
                if status in ("done", "failed"):
                    task.completed_at = __import__('datetime').datetime.utcnow()
                if status == "failed":
                    task.error_msg = step
                # 报告摘要（看板加速查询）
                if score is not None:
                    task.score = score
                    task.critical_issues = critical_issues
                    task.major_issues = major_issues
                    task.minor_issues = minor_issues
                await db.commit()
            break
    except Exception as e:
        logger.error(f"更新任务状态失败 [{task_id}]: {e}")


async def _run_review_task(task_id, ref_path, draft_path, task_name, ref_paths=None):
    """在后台执行审查任务（独立 db 连接，不依赖外部传入的 session）"""
    from services.review_service import ReviewService

    async def progress_updater(tid, progress, step, status,
                               parsing_started_at=None, analysis_started_at=None,
                               score=None, critical_issues=0, major_issues=0, minor_issues=0):
        await _update_task_in_db(tid, progress, step, status,
                                 parsing_started_at, analysis_started_at,
                                 score, critical_issues, major_issues, minor_issues)

    service = ReviewService(progress_updater=progress_updater)
    try:
        report_data = await service.run_review(
            task_id, draft_path, ref_path, task_name, ref_paths=ref_paths
        )

        # 保存报告记录
        async for db in get_db():
            report = ReviewReport(
                id=uuid.uuid4().hex,
                task_id=task_id,
                total_issues=report_data["total_issues"],
                critical_issues=report_data["critical_issues"],
                major_issues=report_data["major_issues"],
                minor_issues=report_data["minor_issues"],
                score=report_data["score"],
                result_json=__import__('json').dumps(report_data, ensure_ascii=False)
            )
            db.add(report)
            await db.commit()

            # 生成 HTML 报告
            report_service.save_html_report(report_data)
            report.html_report_path = os.path.join("data", "reports", f"{task_id}.html")
            await db.commit()
            break

        # 把报告摘要同步回 ReviewTask（供看板列表快速查询）
        await _update_task_in_db(
            task_id,
            progress=100,
            step="审查完成",
            status="done",
            score=report_data["score"],
            critical_issues=report_data["critical_issues"],
            major_issues=report_data["major_issues"],
            minor_issues=report_data["minor_issues"]
        )

    except Exception as e:
        logger.error(f"审查任务 {task_id} 失败: {e}", exc_info=True)
    finally:
        await service.close()


@router.get("/reviews")
async def list_reviews():
    """获取审查任务列表（含标准名称和批次信息，供看板使用）"""
    async for db in get_db():
        from sqlalchemy import select
        result = await db.execute(select(ReviewTask).order_by(ReviewTask.created_at.desc()))
        tasks = result.scalars().all()

        # 批量查询标准名称（避免 N+1）
        all_std_ids = set()
        for t in tasks:
            if t.standard_id:
                all_std_ids.update(sid.strip() for sid in t.standard_id.split(",") if sid.strip())

        std_names_map = {}
        if all_std_ids:
            std_result = await db.execute(select(Standard).where(Standard.id.in_(all_std_ids)))
            for s in std_result.scalars().all():
                std_names_map[s.id] = s.name

        # 批量查询批次名称
        batch_ids = set(t.batch_id for t in tasks if t.batch_id)
        batch_names_map = {}
        if batch_ids:
            batch_result = await db.execute(select(ReviewBatch).where(ReviewBatch.id.in_(batch_ids)))
            for b in batch_result.scalars().all():
                batch_names_map[b.id] = b.name

        items = []
        for t in tasks:
            item = ReviewTaskOut.model_validate(t).model_dump()
            # 附加标准名称
            if t.standard_id:
                names = [std_names_map.get(sid.strip(), sid.strip())
                         for sid in t.standard_id.split(",") if sid.strip()]
                item["standard_names"] = " / ".join(names)
            else:
                item["standard_names"] = ""
            # 附加批次名称
            if t.batch_id:
                item["batch_name"] = batch_names_map.get(t.batch_id, t.batch_id)
            else:
                item["batch_name"] = None
            items.append(item)

        return ApiResponse(success=True, data=items)


@router.delete("/reviews/{task_id}")
async def delete_review(task_id: str):
    """删除审查任务及相关报告"""
    import os as _os
    async for db in get_db():
        from sqlalchemy import select
        result = await db.execute(select(ReviewTask).where(ReviewTask.id == task_id))
        task = result.scalar_one_or_none()
        if not task:
            raise HTTPException(status_code=404, detail="审查任务不存在")

        # 删除初稿文件
        if task.draft_file_path and _os.path.exists(task.draft_file_path):
            try:
                _os.remove(task.draft_file_path)
            except Exception as e:
                logger.warning(f"删除初稿文件失败: {e}")

        # 删除 HTML 报告文件
        html_report_path = _os.path.join("data", "reports", f"{task_id}.html")
        if _os.path.exists(html_report_path):
            try:
                _os.remove(html_report_path)
            except Exception as e:
                logger.warning(f"删除报告文件失败: {e}")

        # 删除关联的审查报告记录
        result_report = await db.execute(select(ReviewReport).where(ReviewReport.task_id == task_id))
        report = result_report.scalar_one_or_none()
        if report:
            await db.delete(report)

        # 删除审查任务记录
        await db.delete(task)
        await db.commit()
        return ApiResponse(success=True, message="删除成功")


@router.get("/reviews/{task_id}/status")
async def get_review_status(task_id: str):
    """获取审查任务状态"""
    async for db in get_db():
        from sqlalchemy import select
        result = await db.execute(select(ReviewTask).where(ReviewTask.id == task_id))
        task = result.scalar_one_or_none()
        if not task:
            return ApiResponse(success=False, message="任务不存在")
        return ApiResponse(success=True, data={
            "task_id": task.id,
            "status": task.status,
            "progress": task.progress,
            "current_step": task.current_step,
            "name": task.name,
            "created_at": task.created_at.isoformat() if task.created_at else None,
            "parsing_started_at": task.parsing_started_at.isoformat() if task.parsing_started_at else None,
            "analysis_started_at": task.analysis_started_at.isoformat() if task.analysis_started_at else None,
            "completed_at": task.completed_at.isoformat() if task.completed_at else None,
        })


@router.get("/reviews/{task_id}/report")
async def get_review_report(task_id: str):
    """获取审查报告数据"""
    report = report_service.load_report(task_id)
    if not report:
        # 尝试从数据库加载
        async for db in get_db():
            from sqlalchemy import select
            result = await db.execute(select(ReviewReport).where(ReviewReport.task_id == task_id))
            r = result.scalar_one_or_none()
            if r and r.result_json:
                report = __import__('json').loads(r.result_json)
                break

    if not report:
        return ApiResponse(success=False, message="报告不存在")

    return ApiResponse(success=True, data=report)


@router.get("/reviews/{task_id}/report/html")
async def get_review_report_html(task_id: str):
    """获取审查报告 HTML"""
    html_path = report_service.get_html_report_path(task_id)
    if html_path:
        return FileResponse(html_path, media_type="text/html")

    # 动态生成
    report = report_service.load_report(task_id)
    if not report:
        async for db in get_db():
            from sqlalchemy import select
            result = await db.execute(select(ReviewReport).where(ReviewReport.task_id == task_id))
            r = result.scalar_one_or_none()
            if r and r.result_json:
                report = __import__('json').loads(r.result_json)
                break

    if not report:
        raise HTTPException(status_code=404, detail="报告不存在")

    html = report_service.generate_html_report(report)
    return HTMLResponse(content=html)


@router.get("/reviews/{task_id}/report/download")
async def download_report(task_id: str):
    """下载 HTML 报告文件"""
    html_path = report_service.get_html_report_path(task_id)
    if not html_path:
        # 尝试生成
        report = report_service.load_report(task_id)
        if report:
            html_path = report_service.save_html_report(report)
    if html_path:
        return FileResponse(
            html_path,
            media_type="application/octet-stream",
            filename=f"review_report_{task_id}.html"
        )
    raise HTTPException(status_code=404, detail="报告不存在")


@router.get("/reviews/{task_id}/report/docx")
async def download_report_docx(task_id: str):
    """下载 Word 报告文件"""
    # 获取报告数据
    report = report_service.load_report(task_id)
    if not report:
        async for db in get_db():
            from sqlalchemy import select
            result = await db.execute(select(ReviewReport).where(ReviewReport.task_id == task_id))
            r = result.scalar_one_or_none()
            if r and r.result_json:
                report = __import__('json').loads(r.result_json)
                break

    if not report:
        raise HTTPException(status_code=404, detail="报告不存在")

    # 获取任务名称用于文件名
    async for db in get_db():
        from sqlalchemy import select
        result = await db.execute(select(ReviewTask).where(ReviewTask.id == task_id))
        task = result.scalar_one_or_none()
        task_name = task.name if task else task_id
        break

    # 生成或读取 docx
    docx_path = report_service.get_docx_report_path(task_id)
    if not docx_path:
        docx_path = report_service.save_docx_report(report)

    # 清理任务名称中的非法文件名字符
    safe_name = "".join(c for c in task_name if c.isalnum() or c in (' ', '-', '_', '（', '）')).strip()
    filename = f"审查报告_{safe_name}_{task_id[:8]}.docx"

    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )
