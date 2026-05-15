import json
import os
from datetime import datetime
from typing import Optional
from jinja2 import Template

import sys, os as _os
sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
from config import REPORTS_DIR

# 内联 HTML 报告模板
HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>标准审查报告 - {{ task_name }}</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: 'Noto Sans SC', -apple-system, sans-serif; background: #f3f4f6; color: #111827; line-height: 1.6; }
        .container { max-width: 960px; margin: 0 auto; padding: 24px; }
        .report-header { background: linear-gradient(135deg, #1A56DB 0%, #1E6FF2 100%); color: white; padding: 32px; border-radius: 12px; margin-bottom: 24px; }
        .report-header h1 { font-size: 24px; font-weight: 600; margin-bottom: 8px; }
        .report-header p { opacity: 0.85; font-size: 14px; }
        .score-card { display: inline-flex; align-items: center; justify-content: center; width: 80px; height: 80px; border-radius: 50%; background: white; color: #1A56DB; font-size: 28px; font-weight: 700; margin-top: 16px; }
        .score-card.excellent { color: #10B981; }
        .score-card.good { color: #F59E0B; }
        .score-card.poor { color: #DC2626; }
        .stats-grid { display: grid; grid-template-columns: repeat(4, 1fr); gap: 16px; margin-bottom: 24px; }
        .stat-card { background: white; padding: 20px; border-radius: 10px; text-align: center; box-shadow: 0 1px 3px rgba(0,0,0,0.08); }
        .stat-card .num { font-size: 32px; font-weight: 700; }
        .stat-card .label { font-size: 13px; color: #6B7280; margin-top: 4px; }
        .stat-card.critical .num { color: #DC2626; }
        .stat-card.major .num { color: #F59E0B; }
        .stat-card.minor .num { color: #3B82F6; }
        .stat-card.score .num { color: #10B981; }
        .meta-row { display: flex; gap: 24px; margin-bottom: 20px; flex-wrap: wrap; }
        .meta-item { font-size: 13px; color: #6B7280; }
        .meta-item strong { color: #374151; }
        .filter-bar { display: flex; gap: 8px; margin-bottom: 20px; flex-wrap: wrap; }
        .filter-btn { padding: 6px 16px; border-radius: 20px; border: 1px solid #d1d5db; background: white; cursor: pointer; font-size: 13px; transition: all 0.2s; }
        .filter-btn.active { background: #1A56DB; color: white; border-color: #1A56DB; }
        .filter-btn:hover { border-color: #1A56DB; }
        .issue-card { background: white; border-radius: 10px; padding: 16px; margin-bottom: 12px; box-shadow: 0 1px 3px rgba(0,0,0,0.06); border-left: 4px solid #d1d5db; }
        .issue-card.critical { border-left-color: #DC2626; }
        .issue-card.major { border-left-color: #F59E0B; }
        .issue-card.minor { border-left-color: #3B82F6; }
        .issue-card.suggestion { border-left-color: #10B981; }
        .issue-header { display: flex; align-items: center; gap: 8px; margin-bottom: 8px; }
        .issue-badge { font-size: 11px; padding: 2px 8px; border-radius: 4px; font-weight: 500; }
        .badge-critical { background: #FEE2E2; color: #DC2626; }
        .badge-major { background: #FEF3C7; color: #D97706; }
        .badge-minor { background: #DBEAFE; color: #2563EB; }
        .badge-suggestion { background: #D1FAE5; color: #059669; }
        .category-badge { font-size: 11px; padding: 2px 8px; border-radius: 4px; background: #f3f4f6; color: #6B7280; }
        .section-tag { font-size: 12px; color: #6B7280; }
        .issue-title { font-size: 15px; font-weight: 500; color: #111827; }
        .issue-body { font-size: 13px; color: #374151; margin-top: 8px; }
        .issue-ref { font-size: 12px; color: #6B7280; margin-top: 6px; font-style: italic; }
        .issue-suggestion { margin-top: 8px; padding: 10px; background: #f0fdf4; border-radius: 6px; font-size: 13px; color: #065f46; }
        .issue-suggestion strong { color: #047857; }
        .empty-state { text-align: center; padding: 40px; color: #9CA3AF; }
        @media (max-width: 640px) {
            .stats-grid { grid-template-columns: repeat(2, 1fr); }
            .report-header { padding: 20px; }
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="report-header">
            <h1>标准审查报告</h1>
            <p>{{ task_name }} · {{ created_at }}</p>
            <div class="score-card {{ score_class }}">{{ score }}</div>
        </div>

        <div class="meta-row">
            <div class="meta-item"><strong>待审初稿：</strong>{{ draft_title }}</div>
            <div class="meta-item"><strong>参考标准：</strong>{{ reference_title }}</div>
        </div>

        <div class="stats-grid">
            <div class="stat-card score"><div class="num">{{ score }}</div><div class="label">综合评分</div></div>
            <div class="stat-card critical"><div class="num">{{ critical_issues }}</div><div class="label">严重问题</div></div>
            <div class="stat-card major"><div class="num">{{ major_issues }}</div><div class="label">一般问题</div></div>
            <div class="stat-card minor"><div class="num">{{ minor_issues }}</div><div class="label">轻微问题</div></div>
        </div>

        <div class="filter-bar">
            <button class="filter-btn active" onclick="filterIssues('all')">全部 ({{ total_issues }})</button>
            <button class="filter-btn" onclick="filterIssues('critical')">严重 ({{ critical_issues }})</button>
            <button class="filter-btn" onclick="filterIssues('major')">一般 ({{ major_issues }})</button>
            <button class="filter-btn" onclick="filterIssues('minor')">轻微 ({{ minor_issues }})</button>
            <button class="filter-btn" onclick="filterIssues('format')">格式合规</button>
            <button class="filter-btn" onclick="filterIssues('completeness')">内容完整性</button>
            <button class="filter-btn" onclick="filterIssues('terminology')">术语一致性</button>
            <button class="filter-btn" onclick="filterIssues('semantic')">语义对比</button>
        </div>

        <div id="issues-list">
            {% for issue in issues %}
            <div class="issue-card {{ issue.level }}" data-level="{{ issue.level }}" data-category="{{ issue.category }}">
                <div class="issue-header">
                    <span class="issue-badge badge-{{ issue.level }}">{{ level_labels[issue.level] }}</span>
                    <span class="category-badge">{{ category_labels[issue.category] }}</span>
                    {% if issue.section and issue.section != '-' %}
                    <span class="section-tag">{{ issue.section }}</span>
                    {% endif %}
                    <span class="issue-title">{{ issue.title }}</span>
                </div>
                <div class="issue-body">{{ issue.description }}</div>
                {% if issue.reference %}
                <div class="issue-ref">参考依据：{{ issue.reference }}</div>
                {% endif %}
                {% if issue.suggestion %}
                <div class="issue-suggestion"><strong>改进建议：</strong>{{ issue.suggestion }}</div>
                {% endif %}
            </div>
            {% endfor %}
            {% if not issues %}
            <div class="empty-state">未发现问题，文档质量良好。</div>
            {% endif %}
        </div>
    </div>

    <script>
    const levelLabels = { critical: '严重', major: '一般', minor: '轻微', suggestion: '建议' };
    const categoryLabels = { format: '格式合规', completeness: '内容完整性', terminology: '术语一致性', semantic: '语义对比' };
    function filterIssues(type) {
        document.querySelectorAll('.filter-btn').forEach(b => b.classList.remove('active'));
        event.target.classList.add('active');
        document.querySelectorAll('.issue-card').forEach(card => {
            if (type === 'all') { card.style.display = ''; return; }
            const level = card.dataset.level;
            const cat = card.dataset.category;
            card.style.display = (level === type || cat === type) ? '' : 'none';
        });
    }
    </script>
</body>
</html>"""


class ReportService:
    """报告生成服务"""

    def __init__(self):
        self.template = Template(HTML_TEMPLATE)

    def generate_html_report(self, report_data: dict) -> str:
        """生成 HTML 报告字符串"""
        report_data["level_labels"] = {
            "critical": "严重", "major": "一般", "minor": "轻微", "suggestion": "建议"
        }
        report_data["category_labels"] = {
            "format": "格式合规", "completeness": "内容完整性",
            "terminology": "术语一致性", "semantic": "语义对比"
        }
        # 评分等级
        score = report_data.get("score", 100)
        if score >= 80:
            report_data["score_class"] = "excellent"
        elif score >= 60:
            report_data["score_class"] = "good"
        else:
            report_data["score_class"] = "poor"

        # 格式化日期
        created = report_data.get("created_at", "")
        if created:
            try:
                dt = datetime.fromisoformat(created)
                report_data["created_at"] = dt.strftime("%Y-%m-%d %H:%M")
            except Exception:
                pass

        return self.template.render(**report_data)

    def save_html_report(self, report_data: dict) -> str:
        """保存 HTML 报告到文件，返回路径"""
        html = self.generate_html_report(report_data)
        task_id = report_data.get("task_id", "unknown")
        path = os.path.join(REPORTS_DIR, f"{task_id}.html")
        with open(path, "w", encoding="utf-8") as f:
            f.write(html)
        return path

    def load_report(self, task_id: str) -> Optional[dict]:
        """加载审查报告数据"""
        path = os.path.join(REPORTS_DIR, f"{task_id}.json")
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        return None

    def get_html_report_path(self, task_id: str) -> Optional[str]:
        """获取 HTML 报告文件路径"""
        path = os.path.join(REPORTS_DIR, f"{task_id}.html")
        if os.path.exists(path):
            return path
        return None

    def _make_docx(self, report_data: dict) -> 'Document':
        """根据报告数据生成 Word 文档（python-docx Document 对象）"""
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # 页面边距
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        LEVEL_COLORS = {
            "critical": RGBColor(0xDC, 0x26, 0x26),
            "major":    RGBColor(0xF5, 0x9E, 0x0B),
            "minor":    RGBColor(0x25, 0x63, 0xEB),
            "suggestion": RGBColor(0x10, 0xB9, 0x81),
        }
        LEVEL_LABELS = {"critical": "严重", "major": "一般", "minor": "轻微", "suggestion": "建议"}
        CAT_LABELS = {
            "format": "格式合规", "completeness": "内容完整性",
            "terminology": "术语一致性", "semantic": "语义对比"
        }

        # 标题
        title = doc.add_heading("标准审查报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

        # 副标题
        subtitle_text = f"{report_data.get('task_name', '')}  ·  {report_data.get('created_at', '')}"
        sub = doc.add_paragraph(subtitle_text)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        doc.add_paragraph()  # 空行

        # --- 汇总信息表 ---
        doc.add_heading("审查概览", level=1)
        score = report_data.get("score", 100)
        critical = report_data.get("critical_issues", 0)
        major = report_data.get("major_issues", 0)
        minor = report_data.get("minor_issues", 0)
        total = report_data.get("total_issues", 0)

        tbl = doc.add_table(rows=5, cols=2)
        tbl.style = "Table Grid"
        rows_data = [
            ("综合评分", f"{score} 分"),
            ("严重问题", str(critical)),
            ("一般问题", str(major)),
            ("轻微问题", str(minor)),
            ("问题合计", str(total)),
        ]
        for i, (k, v) in enumerate(rows_data):
            tbl.rows[i].cells[0].text = k
            tbl.rows[i].cells[1].text = v
            # 评分行高亮
            if k == "综合评分":
                for cell in tbl.rows[i].cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if score >= 80:
                                run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
                            elif score >= 60:
                                run.font.color.rgb = RGBColor(0xF5, 0x9E, 0x0B)
                            else:
                                run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                            run.font.bold = True

        doc.add_paragraph()

        # --- 基本信息 ---
        doc.add_heading("基本信息", level=1)
        meta_tbl = doc.add_table(rows=2, cols=2)
        meta_tbl.style = "Table Grid"
        meta_tbl.rows[0].cells[0].text = "待审初稿"
        meta_tbl.rows[0].cells[1].text = report_data.get("draft_title", "—")
        meta_tbl.rows[1].cells[0].text = "参考标准"
        meta_tbl.rows[1].cells[1].text = report_data.get("reference_title", "—")

        doc.add_paragraph()

        # --- 问题详情 ---
        issues = report_data.get("issues", [])
        if issues:
            doc.add_heading(f"问题详情（共 {len(issues)} 项）", level=1)

            # 按严重程度分组
            grouped = {}
            for iss in issues:
                lvl = iss.get("level", "minor")
                grouped.setdefault(lvl, []).append(iss)

            for level in ["critical", "major", "minor", "suggestion"]:
                if level not in grouped:
                    continue
                level_issues = grouped[level]
                color = LEVEL_COLORS.get(level, RGBColor(0, 0, 0))
                lbl = LEVEL_LABELS.get(level, level)

                # 级别标题
                h = doc.add_heading(f"【{lbl}】{len(level_issues)} 项", level=2)
                for run in h.runs:
                    run.font.color.rgb = color

                for iss in level_issues:
                    cat = CAT_LABELS.get(iss.get("category", ""), iss.get("category", ""))
                    section_ref = iss.get("section", "") or ""

                    # 问题标题
                    title_para = doc.add_paragraph()
                    title_para.paragraph_format.left_indent = Cm(0.3)
                    run = title_para.add_run(f"{lbl}：{iss.get('title', '')}")
                    run.font.bold = True
                    run.font.color.rgb = color
                    run.font.size = Pt(11)

                    if cat or section_ref:
                        meta_run = title_para.add_run(f"  [{cat}{' · ' + section_ref if section_ref else ''}]")
                        meta_run.font.size = Pt(9)
                        meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

                    # 问题描述
                    desc_para = doc.add_paragraph()
                    desc_para.paragraph_format.left_indent = Cm(0.6)
                    desc_run = desc_para.add_run(f"问题：{iss.get('description', '')}")
                    desc_run.font.size = Pt(10)

                    # 参考依据
                    if iss.get("reference"):
                        ref_para = doc.add_paragraph()
                        ref_para.paragraph_format.left_indent = Cm(0.6)
                        ref_run = ref_para.add_run(f"参考依据：{iss.get('reference', '')}")
                        ref_run.font.size = Pt(9)
                        ref_run.font.italic = True
                        ref_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

                    # 改进建议
                    if iss.get("suggestion"):
                        sug_para = doc.add_paragraph()
                        sug_para.paragraph_format.left_indent = Cm(0.6)
                        sug_para.paragraph_format.space_before = Pt(4)
                        # 浅绿背景通过 XML 实现，这里简化处理
                        sug_run = sug_para.add_run(f"改进建议：{iss.get('suggestion', '')}")
                        sug_run.font.size = Pt(10)
                        sug_run.font.color.rgb = RGBColor(0x04, 0x7F, 0x46)

                    doc.add_paragraph()  # 项间距
        else:
            doc.add_heading("问题详情", level=1)
            doc.add_paragraph("✅ 未发现问题，文档质量良好。")

        # 页脚
        from docx.oxml import OxmlElement
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(f"标准审查助手 · {report_data.get('created_at', '')}")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0x9C, 0xAF, 0xC0)

        return doc

    def save_docx_report(self, report_data: dict) -> str:
        """生成并保存 Word 报告，返回文件路径"""
        from docx import Document
        task_id = report_data.get("task_id", "unknown")
        path = os.path.join(REPORTS_DIR, f"{task_id}.docx")
        doc = self._make_docx(report_data)
        doc.save(path)
        return path

    def get_docx_report_path(self, task_id: str) -> Optional[str]:
        """获取 Word 报告文件路径"""
        path = os.path.join(REPORTS_DIR, f"{task_id}.docx")
        if os.path.exists(path):
            return path
        return None
